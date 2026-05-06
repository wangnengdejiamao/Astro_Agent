import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)


INVISIBLE_CHARS_RE = re.compile(r'[\u200b\u200c\u200d\ufeff\u00ad]')
FORMULA_CHAR_RE = re.compile(
    r'[0-9+\-\u2212\u2192\u2190\u22c5\u00b7()@/=\u02da\u00c5'
    r'\u00b2\u00b3\u00b9\u2082\u2083\u2084\u03bb\u2206\u2207]'
)

SHORT_THRESHOLD = 3
MEDIUM_THRESHOLD = 12
GAP_TOLERANCE = 2
MIN_FRAGMENT_COUNT = 2


def _clean(text):
    return INVISIBLE_CHARS_RE.sub('', text.strip())


def _extract_and_merge(doc):
    """Extract paragraph text from Document, merge formula fragments."""
    all_paras = []
    for para in doc.paragraphs:
        raw = para.text.strip()
        cleaned = _clean(raw)
        all_paras.append({
            'raw': raw,
            'cleaned': cleaned,
            'clen': len(cleaned),
        })

    n = len(all_paras)
    is_fragment = [False] * n

    i = 0
    while i < n:
        p = all_paras[i]
        if p['clen'] == 0 or p['clen'] > SHORT_THRESHOLD:
            i += 1
            continue

        seq_start = i
        j = i
        gap_count = 0

        while j < n:
            pj = all_paras[j]
            if pj['clen'] == 0:
                gap_count += 1
                if gap_count > GAP_TOLERANCE:
                    break
                j += 1
            elif 0 < pj['clen'] <= SHORT_THRESHOLD:
                gap_count = 0
                j += 1
            elif 0 < pj['clen'] <= MEDIUM_THRESHOLD:
                found_short = False
                for k in range(j + 1, min(j + GAP_TOLERANCE + 2, n)):
                    if 0 < all_paras[k]['clen'] <= SHORT_THRESHOLD:
                        found_short = True
                        break
                if found_short:
                    gap_count = 0
                    j += 1
                else:
                    break
            else:
                break

        seq_end = j
        non_empty = sum(1 for k in range(seq_start, seq_end) if all_paras[k]['clen'] > 0)

        if non_empty >= MIN_FRAGMENT_COUNT:
            for k in range(seq_start, seq_end):
                is_fragment[k] = True

        i = max(j, i + 1)

    merged = []
    i = 0
    while i < n:
        if is_fragment[i]:
            fragments = []
            while i < n and is_fragment[i]:
                if all_paras[i]['clen'] > 0:
                    fragments.append(all_paras[i]['cleaned'])
                i += 1

            formula_text = "".join(fragments)
            if formula_text and merged:
                merged[-1] = merged[-1] + formula_text
            elif formula_text:
                merged.append(formula_text)
        else:
            if all_paras[i]['clen'] > 0:
                merged.append(all_paras[i]['raw'])
            i += 1

    return merged


def _dedup_within_line(line):
    """Remove adjacent duplicate substrings within a single line."""
    changed = True
    while changed:
        changed = False
        for length in range(min(30, len(line) // 2), 2, -1):
            pattern = re.compile('(.{' + str(length) + '})\\1')
            m = pattern.search(line)
            if m:
                chunk = m.group(1)
                has_special = bool(FORMULA_CHAR_RE.search(chunk))
                is_short_formula = len(chunk) <= 10 and not chunk.isalpha()
                if has_special or is_short_formula:
                    line = line[:m.start() + length] + line[m.end():]
                    changed = True
                    break
    return line


def _remove_adjacent_duplicates(text):
    """
    Remove adjacent duplicate formula renderings, including across newlines.

    Three-pass approach:
    1. Join short formula-residue lines onto the previous line so duplicates
       that span a newline become within-line duplicates.
    2. Run within-line dedup on every line.
    3. Cross-line overlap removal for any remaining cases.
    """
    lines = text.split('\n')

    # Pass 1: join short formula-like lines onto the previous line
    joined = []
    for line in lines:
        stripped = line.strip()
        if (joined
                and stripped
                and len(stripped) < 20
                and bool(FORMULA_CHAR_RE.search(stripped))
                and not stripped.endswith('.')
                and not (len(stripped) > 3 and stripped[0].isupper() and ' ' in stripped)):
            joined[-1] = joined[-1] + stripped
        else:
            joined.append(line)

    # Pass 2: within-line dedup
    lines = [_dedup_within_line(line) for line in joined]

    # Pass 3: cross-line overlap removal
    result = []
    for line in lines:
        if not result:
            result.append(line)
            continue

        prev = result[-1]
        curr = line.strip()

        if not curr:
            result.append(line)
            continue

        # curr entirely duplicates end of prev
        if len(curr) >= 2 and prev.endswith(curr):
            continue

        # curr starts with a prefix that duplicates the end of prev
        best_overlap = 0
        max_check = min(len(prev), len(curr), 30)
        for length in range(max_check, 1, -1):
            if prev.endswith(curr[:length]):
                overlap_text = curr[:length]
                has_special = bool(FORMULA_CHAR_RE.search(overlap_text))
                if has_special or (len(overlap_text) <= 10 and not overlap_text.isalpha()):
                    best_overlap = length
                    break

        if best_overlap > 0:
            remainder = curr[best_overlap:]
            if remainder.strip():
                result.append(remainder)
        else:
            result.append(line)

    return '\n'.join(result)


def clean_text(text):
    """Clean text: invisible chars, LaTeX, duplicates, whitespace."""
    text = INVISIBLE_CHARS_RE.sub('', text)

    # Handle $$ ... $$ LaTeX blocks
    lines = text.split('\n')
    cleaned_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line == '$$' and i + 2 < len(lines) and lines[i + 2].strip() == '$$':
            formula = lines[i + 1].strip()
            if cleaned_lines:
                cleaned_lines[-1] = cleaned_lines[-1] + ' [' + formula + ']'
            else:
                cleaned_lines.append('[' + formula + ']')
            i += 3
        elif line == '$$':
            i += 1
        else:
            cleaned_lines.append(line)
            i += 1

    text = '\n'.join(cleaned_lines)
    text = _remove_adjacent_duplicates(text)

    # Normalize whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    lines = text.split('\n')
    text = '\n'.join(l for l in lines if l.strip() or l == '')
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def extract_text_from_docx(docx_path: str) -> dict:
    """从 DOCX 文件提取 title 和 text，自动处理公式碎片"""
    doc = Document(docx_path)

    merged_texts = _extract_and_merge(doc)
    text_content = "\n".join(merged_texts)
    text_content = clean_text(text_content)

    title = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            if paragraph.style.name.startswith('Heading'):
                title = text
                break
            elif len(text) > 5 and len(text) < 200:
                title = text
                break

    if not title:
        title = Path(docx_path).stem

    title = INVISIBLE_CHARS_RE.sub('', title)

    return {
        "title": title,
        "text": text_content
    }


def docx_to_corpus_json(docx_path: str, output_dir: str = None) -> str:
    """将 DOCX 文件转换为 corpus_cleaned.json 格式"""
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")

    corpus_item = extract_text_from_docx(docx_path)
    print(f"提取标题: {corpus_item['title']}")
    print(f"文本长度: {len(corpus_item['text'])} 字符")

    if output_dir is None:
        corpus_name = docx_file.stem
        project_dir = os.getenv("PROJECT_DIR") or os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(project_dir, "input", corpus_name)

    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "corpus_cleaned.json")

    corpus = [corpus_item]
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(corpus, f, ensure_ascii=False, indent=2)

    print(f"\n已保存到: {output_path}")
    return output_path


if __name__ == "__main__":
    docx_path = "/devSpaceIT/huangjiahao/prompt2graph/process_docx/Paper-456.docx"

    output_path = docx_to_corpus_json(docx_path)
    print(f"\n完成！可以使用此 corpus 运行 pipeline:")
    print(f"  corpus_path: input/{Path(docx_path).stem}/corpus_cleaned.json")
