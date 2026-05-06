"""
Document Parser Utility
Supports parsing PDF, DOCX, DOC files using MinerU and python-docx
"""

import os
import tempfile
from typing import Optional, Dict
from pathlib import Path

from utils.logger import logger

try:
    from magic_pdf.data.dataset import PymuDocDataset
    MINERU_AVAILABLE = True
except ImportError as e:
    MINERU_AVAILABLE = False
    logger.warning(f"MinerU not available: {e}")

try:
    import fitz  # type: ignore[attr-defined]
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) not available; PDF text extraction may be limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

try:
    import subprocess
    # Check for system-level antiword (prefer /usr/local/bin over pip package)
    antiword_check = subprocess.run(['which', '/usr/local/bin/antiword'], 
                                    capture_output=True)
    if antiword_check.returncode == 0:
        ANTIWORD_PATH = '/usr/local/bin/antiword'
        ANTIWORD_AVAILABLE = True
    else:
        # Fallback to PATH search
        antiword_check = subprocess.run(['which', 'antiword'], 
                                       capture_output=True)
        ANTIWORD_PATH = 'antiword'
        ANTIWORD_AVAILABLE = antiword_check.returncode == 0
except Exception:
    ANTIWORD_AVAILABLE = False
    ANTIWORD_PATH = 'antiword'

try:
    import textract  # type: ignore
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    logger.debug("textract not available")

try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.debug("Apache Tika not available")

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False
    logger.debug("striprtf not available - RTF file support limited")


class DocumentParser:
    """Parse various document formats to extract text content"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="youtu_graphrag_")
        logger.info(f"DocumentParser initialized with temp dir: {self.temp_dir}")
    
    def parse_file(self, file_path: str, file_type: str) -> Optional[str]:
        """
        Parse a document file and extract text content
        
        Args:
            file_path: Path to the document file
            file_type: File extension (.pdf, .docx, .doc)
            
        Returns:
            Extracted text content or None if parsing fails
        """
        file_type = file_type.lower()
        
        try:
            if file_type == '.pdf':
                return self._parse_pdf(file_path)
            elif file_type in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error parsing {file_type} file: {e}")
            return None
    
    def _parse_pdf(self, pdf_path: str) -> Optional[str]:
        """
        Parse PDF using MinerU
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not MINERU_AVAILABLE:
            logger.error("MinerU is not installed. Cannot parse PDF files.")
            return None
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
        except Exception as e:
            logger.error(f"Unable to read PDF file {pdf_path}: {e}")
            return None

        # -------- MinerU / PyMuPDF pipeline --------
        if MINERU_AVAILABLE:
            try:
                dataset = PymuDocDataset(pdf_bytes, lang='auto')
                text_parts: list[str] = []
                for page_index in range(len(dataset)):
                    try:
                        page_doc = dataset.get_page(page_index).get_doc()
                        page_text = page_doc.get_text("text")
                        if page_text:
                            text_parts.append(page_text.strip())
                    except Exception as page_err:
                        logger.warning(
                            f"MinerU pipeline failed to extract page {page_index} of {pdf_path}: {page_err}"
                        )
                        continue

                if text_parts:
                    extracted_text = '\n\n'.join(text_parts)
                    logger.info(
                        f"Successfully extracted {len(extracted_text)} chars from PDF via MinerU pipeline"
                    )
                    return extracted_text
            except Exception as e:
                logger.warning(f"MinerU PDF parsing failed: {e}")

        # -------- Direct PyMuPDF fallback --------
        if PYMUPDF_AVAILABLE:
            try:
                text_parts: list[str] = []
                with fitz.open(pdf_path) as doc:  # type: ignore[attr-defined]
                    for page in doc:
                        page_text = page.get_text("text")
                        if page_text:
                            text_parts.append(page_text.strip())
                if text_parts:
                    extracted_text = '\n\n'.join(text_parts)
                    logger.info(
                        f"Successfully extracted {len(extracted_text)} chars from PDF via PyMuPDF fallback"
                    )
                    return extracted_text
            except Exception as e:
                logger.error(f"PyMuPDF fallback failed: {e}")

        # -------- PyPDF fallback --------
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(pdf_path)
            text_parts = []
            for page_index, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                except Exception as page_err:
                    logger.warning(
                        f"pypdf failed to extract page {page_index} of {pdf_path}: {page_err}"
                    )
                    continue
                if page_text:
                    text_parts.append(page_text.strip())

            if text_parts:
                extracted_text = '\n\n'.join(text_parts)
                logger.info(
                    f"Successfully extracted {len(extracted_text)} chars from PDF via pypdf fallback"
                )
                return extracted_text
        except ImportError:
            logger.debug("pypdf not installed; skipping pypdf fallback")
        except Exception as e:
            logger.error(f"pypdf fallback failed: {e}")

        logger.error(f"Unable to extract text from PDF: {pdf_path}")
        return None
    
    def _parse_docx(self, docx_path: str) -> Optional[str]:
        """
        Parse DOCX/DOC using available methods
        
        Args:
            docx_path: Path to DOCX/DOC file
            
        Returns:
            Extracted text content
        """
        file_ext = os.path.splitext(docx_path)[1].lower()
        
        # Check if the file is actually RTF (some .doc files are RTF in disguise)
        if self._is_rtf_file(docx_path):
            logger.info(f"Detected RTF file disguised as {file_ext}: {docx_path}")
            result = self._parse_rtf(docx_path)
            if result:
                return result
        
        # Try python-docx first for .docx files
        if file_ext == '.docx' and DOCX_AVAILABLE:
            result = self._parse_with_python_docx(docx_path)
            if result:
                return result
        
        # For .doc files or if python-docx fails, try alternative methods
        if file_ext == '.doc':
            # Priority 1: Try antiword first (fast, stable, no Python dependencies)
            if ANTIWORD_AVAILABLE:
                result = self._parse_with_antiword(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via antiword")
                    return result
            
            # Priority 2: Try Apache Tika (best for WPS and complex formats)
            if TIKA_AVAILABLE:
                result = self._parse_with_tika(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via Apache Tika")
                    return result
            
            # Priority 3: Try textract (if available, but has pip 24.1+ conflicts)
            if TEXTRACT_AVAILABLE:
                result = self._parse_with_textract(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via textract")
                    return result
            
            # Priority 4: Try LibreOffice conversion (best for WPS/legacy formats)
            logger.debug(f"Trying LibreOffice for .doc file: {docx_path}")
            result = self._parse_doc_with_libreoffice(docx_path)
            if result:
                logger.info(f"Successfully extracted {len(result)} chars from DOC via LibreOffice")
                return result
            else:
                logger.debug(f"LibreOffice parsing returned None for: {docx_path}")
        
        # Final fallback: try python-docx anyway (might work for some .doc files)
        if DOCX_AVAILABLE:
            result = self._parse_with_python_docx(docx_path)
            if result:
                return result
        
        # Check file type to provide better error message
        file_type_hint = ""
        is_corrupted = False
        try:
            import subprocess
            file_info = subprocess.run(['file', docx_path], capture_output=True, text=True)
            if file_info.returncode == 0:
                info_lower = file_info.stdout.lower()
                if 'wps' in info_lower:
                    file_type_hint = " (WPS Office document)"
                    is_corrupted = True  # WPS 文档可能无法被 LibreOffice 解析
                elif 'composite document' in info_lower or 'ole' in info_lower:
                    if 'microsoft' in info_lower:
                        file_type_hint = " (Legacy Microsoft Word format)"
                    else:
                        file_type_hint = " (OLE document)"
        except Exception as e:
            logger.warning(f"Failed to determine file type for {docx_path}: {e}")
        
        logger.error(f"Unable to parse {file_ext} file{file_type_hint}: {docx_path}")
        
        if is_corrupted:
            logger.warning("⚠️  This WPS Office document cannot be parsed by available tools")
            logger.info("Recommended solutions:")
            logger.info("  1. Open in WPS Office and save as .docx format")
            logger.info("  2. Open in Microsoft Word and save as .docx format")
            logger.info("  3. Use online converters (e.g., zamzar.com, cloudconvert.com)")
        else:
            logger.info("Recommended solutions:")
            logger.info("  1. Install LibreOffice (best compatibility): sudo yum install libreoffice-headless")
            logger.info("  2. Convert file to .docx format using Microsoft Word or WPS")
            logger.info("  3. For standard .doc files: Install antiword (sudo apt-get install antiword)")
        return None
    
    def _parse_with_python_docx(self, docx_path: str) -> Optional[str]:
        """Parse using python-docx library"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = DocxDocument(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                return None
            
            logger.info(f"Successfully extracted {len(extracted_text)} chars via python-docx")
            return extracted_text
            
        except Exception as e:
            logger.debug(f"python-docx failed: {e}")
            return None
    
    def _parse_with_textract(self, doc_path: str) -> Optional[str]:
        """Parse using textract library"""
        if not TEXTRACT_AVAILABLE:
            return None
        
        try:
            text = textract.process(doc_path).decode('utf-8')
            if text and text.strip():
                return text.strip()
        except Exception as e:
            logger.debug(f"textract failed: {e}")
        return None
    
    def _parse_with_tika(self, doc_path: str) -> Optional[str]:
        """
        Parse using Apache Tika (supports WPS, legacy Word, and many other formats)
        
        Args:
            doc_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        if not TIKA_AVAILABLE:
            return None
        
        try:
            # Apache Tika 可以处理几乎所有文档格式,包括 WPS Office 文档
            parsed = tika_parser.from_file(doc_path)
            text = parsed.get('content', '')
            
            if text and text.strip():
                return text.strip()
            else:
                logger.debug(f"Tika returned empty content for: {doc_path}")
                return None
                
        except Exception as e:
            logger.debug(f"Apache Tika parsing failed: {e}")
            return None
    
    def _parse_with_antiword(self, doc_path: str) -> Optional[str]:
        """Parse using antiword command-line tool"""
        if not ANTIWORD_AVAILABLE:
            return None
        
        try:
            result = subprocess.run(
                [ANTIWORD_PATH, doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception as e:
            logger.debug(f"antiword failed: {e}")
        return None
    
    def _parse_doc_with_libreoffice(self, doc_path: str) -> Optional[str]:
        """Convert .doc to .txt using LibreOffice and read the result"""
        try:
            # Check if libreoffice is available
            lo_check = subprocess.run(
                ['which', 'libreoffice'],
                capture_output=True
            )
            if lo_check.returncode != 0:
                logger.debug("LibreOffice not found in PATH")
                return None
            
            logger.debug(f"LibreOffice found, attempting conversion for: {doc_path}")
            
            # Create temp directory for conversion
            import shutil
            temp_dir = tempfile.mkdtemp(prefix="doc_convert_")
            
            try:
                # Copy file to temp dir with safe filename (避免特殊字符导致转换失败)
                import hashlib
                file_ext = os.path.splitext(doc_path)[1]
                # 使用文件路径的哈希生成安全的文件名
                safe_name = hashlib.md5(doc_path.encode()).hexdigest()
                temp_doc = os.path.join(temp_dir, f"{safe_name}{file_ext}")
                shutil.copy2(doc_path, temp_doc)
                
                logger.debug(f"Copied to temp file: {temp_doc}")
                
                # Convert to txt
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', 
                     '--outdir', temp_dir, temp_doc],
                    capture_output=True,
                    text=True,
                    timeout=60,  # 增加超时时间,处理大文件
                    check=False  # 不抛出异常,稍后检查返回码
                )
                
                logger.debug(f"LibreOffice exit code: {result.returncode}")
                if result.stdout:
                    logger.debug(f"LibreOffice stdout: {result.stdout[:200]}")
                if result.stderr:
                    logger.debug(f"LibreOffice stderr: {result.stderr[:200]}")
                
                # Check if conversion succeeded
                if result.returncode != 0:
                    logger.debug(f"LibreOffice conversion failed with code {result.returncode}")
                    return None
                
                # Read the converted file
                txt_path = os.path.join(temp_dir, f"{safe_name}.txt")
                logger.debug(f"Looking for output file: {txt_path}")
                
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    logger.debug(f"Read {len(text)} chars from converted file")
                    if text.strip():
                        return text.strip()
                    else:
                        logger.debug("Converted file is empty")
                        return None
                else:
                    # 即使返回码为0,如果输出文件不存在,说明转换实际失败了
                    logger.debug(f"LibreOffice output file not found: {txt_path}")
                    # 检查stderr中是否有"no export filter"或其他错误
                    if "no export filter" in result.stderr or "Error:" in result.stderr:
                        logger.warning(f"LibreOffice cannot parse this document format. Error: {result.stderr[:150]}")
                    return None
            finally:
                # Cleanup
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                    
        except Exception as e:
            logger.debug(f"LibreOffice conversion exception: {e}")
        return None
    
    def _is_rtf_file(self, file_path: str) -> bool:
        """
        Check if a file is actually RTF format (regardless of extension)
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if file is RTF format
        """
        try:
            with open(file_path, 'rb') as f:
                # RTF files start with {\rtf
                header = f.read(10)
                return header.startswith(b'{\\rtf')
        except Exception as e:
            logger.debug(f"Error checking RTF format: {e}")
            return False
    
    def _parse_rtf(self, rtf_path: str) -> Optional[str]:
        """
        Parse RTF file and extract text content
        
        Args:
            rtf_path: Path to RTF file
            
        Returns:
            Extracted text content
        """
        # Method 1: Try striprtf library if available
        if STRIPRTF_AVAILABLE:
            try:
                with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf_content = f.read()
                text = rtf_to_text(rtf_content)
                if text and text.strip():
                    logger.info(f"Successfully extracted {len(text)} chars from RTF via striprtf")
                    return text.strip()
            except Exception as e:
                logger.debug(f"striprtf parsing failed: {e}")
        
        # Method 2: Try textract if available
        if TEXTRACT_AVAILABLE:
            try:
                text = textract.process(rtf_path).decode('utf-8', errors='ignore')
                if text and text.strip():
                    logger.info(f"Successfully extracted {len(text)} chars from RTF via textract")
                    return text.strip()
            except Exception as e:
                logger.debug(f"textract RTF parsing failed: {e}")
        
        # Method 3: Try LibreOffice conversion
        try:
            # Check if libreoffice is available
            lo_check = subprocess.run(
                ['which', 'libreoffice'],
                capture_output=True
            )
            if lo_check.returncode == 0:
                import hashlib
                import shutil
                temp_dir = tempfile.mkdtemp(prefix="rtf_convert_")
                try:
                    # Copy file to temp dir with safe filename
                    file_ext = os.path.splitext(rtf_path)[1]
                    safe_name = hashlib.md5(rtf_path.encode()).hexdigest()
                    temp_rtf = os.path.join(temp_dir, f"{safe_name}{file_ext}")
                    shutil.copy2(rtf_path, temp_rtf)
                    
                    result = subprocess.run(
                        ['libreoffice', '--headless', '--convert-to', 'txt:Text',
                         '--outdir', temp_dir, temp_rtf],
                        capture_output=True,
                        text=True,
                        timeout=60,
                        check=False
                    )
                    
                    if result.returncode == 0:
                        txt_path = os.path.join(temp_dir, f"{safe_name}.txt")
                        
                        if os.path.exists(txt_path):
                            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                                text = f.read()
                            if text.strip():
                                logger.info(f"Successfully extracted {len(text)} chars from RTF via LibreOffice")
                                return text.strip()
                finally:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
        except Exception as e:
            logger.debug(f"LibreOffice RTF conversion failed: {e}")
        
        logger.error(f"Unable to parse RTF file: {rtf_path}")
        logger.info("Hint: Install 'striprtf' for RTF support: pip install striprtf")
        logger.info("  Or install LibreOffice: sudo yum install libreoffice-headless")
        return None
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp dir: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Error cleaning up temp dir: {e}")


# Global parser instance
_parser_instance = None

def get_parser() -> DocumentParser:
    """Get or create global parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance
